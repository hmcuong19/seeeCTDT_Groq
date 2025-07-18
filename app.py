import streamlit as st
import io
import docx
import fitz  # PyMuPDF
import openai
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# --- Cấu hình và Thiết lập ---
st.set_page_config(page_title="Trích xuất Thông tin Syllabus", page_icon="✨", layout="wide")

# --- API Key cho Groq ---
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except (KeyError, FileNotFoundError):
    st.warning("Không tìm thấy Groq API Key trong Streamlit secrets. Vui lòng nhập thủ công để chạy ứng dụng.")
    GROQ_API_KEY = st.text_input("Nhập Groq API Key của bạn:", type="password")
    if not GROQ_API_KEY:
        st.info("Vui lòng cung cấp API key để bắt đầu.")
        st.stop()

# Khởi tạo client tương thích Groq
client = OpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1"
)

# --- Hàm gọi Groq API ---
def get_groq_response(input_text, prompt, model="llama3-8b-8192"):
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "Bạn là một trợ lý AI thông minh, chuyên trích xuất thông tin từ tài liệu."},
                {"role": "user", "content": input_text},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Đã xảy ra lỗi khi gọi Groq API: {e}"

# --- Hàm xử lý file ---
def extract_text_from_docx(docx_bytes):
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Lỗi đọc file .docx: {e}")
        return None

def extract_text_from_pdf(file_bytes):
    try:
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            full_text += page.get_text()
        pdf_document.close()
        return full_text
    except Exception as e:
        st.error(f"Lỗi đọc file .pdf: {e}")
        return None

# --- Hàm tạo PDF từ nội dung trích xuất ---
def generate_pdf(extracted_text):
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        styles = getSampleStyleSheet()
        
        # Đăng ký font hỗ trợ tiếng Việt
        pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
        
        # Tạo các style tùy chỉnh
        title_style = ParagraphStyle(
            name='Title',
            fontName='DejaVuSans',
            fontSize=16,
            leading=20,
            alignment=1,  # Căn giữa
            spaceAfter=20
        )
        label_style = ParagraphStyle(
            name='Label',
            fontName='DejaVuSans',
            fontSize=10,
            leading=12,
            textColor=colors.darkblue,
            spaceAfter=4
        )
        content_style = ParagraphStyle(
            name='Content',
            fontName='DejaVuSans',
            fontSize=10,
            leading=12,
            spaceAfter=4,
            wordWrap='CJK'  # Hỗ trợ ngắt dòng cho tiếng Việt
        )
        
        # Tạo danh sách các phần tử cho PDF
        story = []
        
        # Tiêu đề
        story.append(Paragraph("Tên học phần", title_style))
        story.append(Spacer(1, 0.5*cm))
        
        # Tách văn bản thành các dòng
        lines = extracted_text.split('\n')
        table_data = []
        current_section = None
        section_content = []
        
        # Phân tích nội dung để tách các mục
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Kiểm tra các mục chính (dựa trên prompt mặc định)
            if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
                if current_section and section_content:
                    # Thêm nội dung của mục trước đó
                    content_text = '\n'.join(section_content)
                    table_data.append([
                        Paragraph(current_section, label_style),
                        Paragraph(content_text, content_style)
                    ])
                current_section = line.split(' ', 1)[1] if ' ' in line else line
                section_content = []
            else:
                section_content.append(line)
        
        # Thêm mục cuối cùng
        if current_section and section_content:
            content_text = '\n'.join(section_content)
            table_data.append([
                Paragraph(current_section, label_style),
                Paragraph(content_text, content_style)
            ])
        
        # Tạo bảng cho các mục chính
        if table_data:
            table = Table(table_data, colWidths=[4*cm, 12.5*cm])
            table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), 'DejaVuSans'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BOX', (0, 0), (-1, -1), 1, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('WORDWRAP', (0, 0), (-1, -1), 'CJK')  # Hỗ trợ ngắt dòng
            ]))
            story.append(table)
        
        # Xây dựng PDF (không có footer)
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Lỗi khi tạo PDF: {e}")
        return None

# --- Giao diện Streamlit ---
st.title("✨ Trích xuất Thông tin từ Tài liệu với Groq AI")
st.markdown("Tải lên tệp `.docx` hoặc `.pdf` để bắt đầu.")

col1, col2 = st.columns([2, 3])

with col1:
    st.header("1. Tải lên & Tùy chỉnh")

    uploaded_file = st.file_uploader("Chọn một tệp (.docx hoặc .pdf)", type=['docx', 'pdf'])

    prompt_default = """Bạn là một trợ lý AI chuyên nghiệp trong việc trích xuất thông tin.

Từ nội dung đề cương học phần cung cấp, hãy trích xuất và trình bày rõ ràng theo kiểu đánh số thứ tự với các mục sau:
Tên học phần
Mã học phần (nếu có)
Số tín chỉ
Điều kiện tiên quyết (nếu có)
Mục tiêu học phần
Chuẩn đầu ra của học phần (CLO)
Nội dung học phần tóm tắt
Tài liệu tham khảo (ghi rõ tên, tác giả, năm, NXB nếu có)

Nếu không tìm thấy thông tin nào, hãy ghi là "Không tìm thấy".
"""
    prompt_user = st.text_area("Chỉnh sửa prompt:", value=prompt_default, height=350)
    submit_button = st.button("🚀 Bắt đầu trích xuất")

with col2:
    st.header("2. Kết quả trích xuất")
    result_container = st.container()
    result_container.info("Kết quả sẽ được hiển thị ở đây sau khi bạn nhấn nút 'Bắt đầu trích xuất'.")

    if submit_button:
        if uploaded_file and prompt_user:
            with st.spinner("Đang xử lý file... Vui lòng chờ! 🤖"):
                file_bytes = uploaded_file.getvalue()
                file_extension = uploaded_file.name.split('.')[-1].lower()
                raw_text = None

                st.info(f"Đang đọc file {file_extension}...")
                if file_extension == "docx":
                    raw_text = extract_text_from_docx(file_bytes)
                elif file_extension == "pdf":
                    raw_text = extract_text_from_pdf(file_bytes)

                if raw_text and raw_text.strip():
                    st.success("Đọc file thành công!")
                    st.info("Đang gửi nội dung đến mô hình AI...")
                    response = get_groq_response(raw_text, prompt_user)
                    result_container.text_area("Thông tin đã trích xuất:", value=response, height=550)
                    
                    # Tạo và cung cấp nút tải PDF
                    pdf_buffer = generate_pdf(response)
                    if pdf_buffer:
                        result_container.download_button(
                            label="📄 Tải xuống kết quả dưới dạng PDF",
                            data=pdf_buffer,
                            file_name="syllabus_information.pdf",
                            mime="application/pdf"
                        )
                    else:
                        result_container.error("Không thể tạo file PDF. Vui lòng thử lại.")
                elif raw_text is not None:
                    result_container.warning("Không tìm thấy nội dung văn bản nào trong file.")
                else:
                    result_container.error("Lỗi khi trích xuất nội dung. Vui lòng thử lại với file khác.")
        elif not uploaded_file:
            st.warning("Vui lòng tải lên một file.")
        else:
            st.warning("Prompt không được để trống.")
